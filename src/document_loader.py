"""Load uploaded files into page-aware text documents."""

from __future__ import annotations

from io import BytesIO
from pathlib import Path
from typing import BinaryIO

from src.models import RawDocument


SUPPORTED_EXTENSIONS = (".pdf", ".docx", ".md", ".txt")


def load_file_bytes(name: str, file_obj: BinaryIO | BytesIO) -> list[RawDocument]:
    extension = Path(name).suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Unsupported file type: {extension}")

    data = file_obj.read()
    if extension == ".pdf":
        return _load_pdf(name, data)
    if extension == ".docx":
        return _load_docx(name, data)

    text = data.decode("utf-8", errors="replace")
    return [RawDocument(source=name, text=clean_text(text))]


def load_local_path(path: str | Path) -> list[RawDocument]:
    local_path = Path(path)
    with local_path.open("rb") as handle:
        return load_file_bytes(local_path.name, handle)


def _load_pdf(name: str, data: bytes) -> list[RawDocument]:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError("PDF support requires pypdf. Install dependencies with: pip install -r requirements.txt") from exc

    reader = PdfReader(BytesIO(data))
    documents: list[RawDocument] = []
    for index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        documents.append(RawDocument(source=name, text=clean_text(text), page=index))
    return documents


def _load_docx(name: str, data: bytes) -> list[RawDocument]:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("DOCX support requires python-docx. Install dependencies with: pip install -r requirements.txt") from exc

    document = Document(BytesIO(data))
    blocks: list[str] = []

    for paragraph in document.paragraphs:
        text = clean_text(paragraph.text)
        if not text:
            continue
        style_name = (paragraph.style.name if paragraph.style else "").lower()
        if style_name.startswith("heading"):
            level = _heading_level(style_name)
            blocks.append(f"{'#' * level} {text}")
        else:
            blocks.append(text)

    for table in document.tables:
        table_lines = _table_lines(table)
        if table_lines:
            blocks.append("\n".join(table_lines))

    return [RawDocument(source=name, text=clean_text("\n\n".join(blocks)))]


def _heading_level(style_name: str) -> int:
    digits = "".join(character for character in style_name if character.isdigit())
    if not digits:
        return 2
    return min(max(int(digits), 1), 6)


def _table_lines(table: object) -> list[str]:
    lines: list[str] = []
    for row in table.rows:
        cells = [clean_text(cell.text) for cell in row.cells]
        cells = [cell for cell in cells if cell]
        if cells:
            lines.append(" | ".join(cells))
    return lines


def clean_text(text: str) -> str:
    lines = [line.strip() for line in text.replace("\x00", " ").splitlines()]
    cleaned: list[str] = []
    previous_blank = False
    for line in lines:
        blank = not line
        if blank and previous_blank:
            continue
        cleaned.append(line)
        previous_blank = blank
    return "\n".join(cleaned).strip()
